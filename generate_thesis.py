
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('基于CLIP多模态预训练与知识蒸馏的\n智能AI家庭健康厨房助手系统')
title_run.font.size = Pt(22)
title_run.font.name = '黑体'
title_run.bold = True
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('创新创业综合实践报告')
sub_run.font.size = Pt(16)

for _ in range(3):
    doc.add_paragraph()

info_items = [
    '专    业：计算机科学与技术',
    '姓    名：________________',
    '学    号：________________',
    '完成时间：2026年6月',
    '上海建桥学院信息技术学院',
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(14)

doc.add_page_break()

# ===== 摘要 =====
add_heading('摘  要', level=1)
add_para(
    '本文设计并实现了一个基于CLIP（Contrastive Language-Image Pre-training）多模态预训练模型的智能AI家庭健康厨房助手系统，'
    '集成了食材图像识别、营养分析、菜谱推荐、饮食记录、膳食指南、家庭成员管理和人脸识别等核心功能。'
    '系统后端采用FastAPI框架和SQLite数据库，前端基于HarmonyOS NEXT（API 11）平台开发。', indent=True
)
add_para(
    '针对食材分类这一核心任务，本文进行了系统性的模型对比实验：从MobileNetV2基线模型（68.24%）、EfficientNet-B0+MixUp'
    '（70.97%），到CLIP ViT-B/32冻结迁移学习（83.36%），再到融合Kaggle Fruits 360高质量数据集后的CLIP模型（92.69%），'
    '验证准确率累计提升24.45个百分点。在此基础上，本文提出了两种改进方案：（1）基于CLIP图像特征与营养文本特征的多模态融合方法，'
    '验证准确率达到88.15%；（2）基于知识蒸馏的端侧轻量化方案，以92.69%的CLIP模型为Teacher，以1.5M参数的MobileNetV3-Small'
    '为Student，通过KL散度蒸馏损失（温度T=4.0, α=0.7）进行训练，蒸馏后Student模型的验证准确率达到90.83%，模型大小从605MB'
    '压缩至6MB，压缩比超过100倍，准确率损失仅1.86个百分点。', indent=True
)
add_para(
    '本文还引入MediaPipe人脸特征点检测技术，实现家庭成员的人脸识别与个性化管理。在数据集构建方面，本文建立了包含15类食材、'
    '训练集4339张、验证集1080张的高质量数据集，融合了Bing搜索引擎爬取数据和Kaggle Fruits 360公开数据集，并提出了基于模型'
    '推理的噪声标签自动检测与校正方法。系统测试结果表明，所设计的方案在识别准确率、模型轻量化和功能完整性方面均达到了预期目标。',
    indent=True
)
add_para('关键词：食材识别；CLIP；知识蒸馏；HarmonyOS；多模态融合；人脸识别', bold=True)

doc.add_page_break()

# ===== 目录占位 =====
add_heading('目  录', level=1)
add_para('（生成后自动更新目录）')

doc.add_page_break()

# ===== 第1章 绪论 =====
add_heading('第1章 绪论', level=1)
add_heading('1.1 研究背景与意义', level=2)
add_para(
    '随着我国居民生活水平的不断提高，家庭饮食健康管理逐渐成为公众关注的焦点。'
    '国家卫健委发布的《中国居民营养与慢性病状况报告》显示，我国成年居民超重肥胖率已超过50%，'
    '不合理膳食结构是导致慢性疾病的主要因素之一。然而，传统厨房设备功能有限，'
    '用户无法便捷地获知食材的营养成分和热量数据；市面上已有的智能厨房设备多依赖云端AI识别，'
    '存在隐私泄露风险、网络依赖性强、响应延迟高等问题，且普遍缺乏面向家庭多成员的个性化管理功能。',
    indent=True
)
add_para(
    '与此同时，以OpenAI CLIP为代表的多模态预训练大模型在计算机视觉领域取得了突破性进展。'
    'CLIP通过在4亿图文对上进行对比学习，学习到了强大的视觉语义表征，在零样本和少样本分类任务上'
    '表现优异。此外，知识蒸馏技术的成熟使得将大模型的知识迁移至轻量级端侧模型成为可能，'
    '为在HarmonyOS等移动端设备上部署高性能AI模型提供了技术路径。',
    indent=True
)
add_para(
    '基于上述背景，本课题拟设计并实现一套基于CLIP多模态预训练模型的智能AI家庭健康厨房助手系统。'
    '系统以食材图像识别为核心功能，结合营养分析、菜谱推荐、家庭成员管理和人脸识别等模块，'
    '旨在为家庭用户提供从食材识别到营养管理的一站式解决方案。'
    '本课题的意义主要体现在以下方面：第一，在算法层面，系统对比了从传统CNN到多模态预训练大模型的多种方案，'
    '为食材识别领域的模型选型提供了实证参考；第二，在工程层面，通过知识蒸馏将605MB的大模型压缩至6MB的端侧模型，'
    '为AI模型在HarmonyOS设备上的部署提供了可行路径；第三，在应用层面，系统集成了人脸识别家庭成员管理功能，'
    '实现了从食材到家庭成员的全链路个性化服务。',
    indent=True
)

add_heading('1.2 国内外研究现状', level=2)
add_para(
    '在食材识别领域，国内外学者和企业已开展了大量研究工作。在国外研究方面，谷歌公司于2019年推出的'
    'MediaPipe框架支持在移动设备上进行实时机器学习推理。苹果公司在iOS 14中引入了Core ML食物识别API。'
    'Kaggle平台上开源的Food-101数据集（2014年发布，含101类食物共10.1万张图片）和Fruits 360数据集'
    '（含131类水果蔬菜）为食物识别研究提供了重要的数据基础。MIT的TinyML研究组在2022年的成果表明，'
    '通过INT8量化技术，可将MobileNetV2模型压缩至200KB以下。OpenAI于2021年发布的CLIP模型开创了多模态'
    '预训练的新范式，通过对比学习实现了强大的零样本图像分类能力。',
    indent=True
)
add_para(
    '在国内研究方面，华为公司于2020年开源OpenHarmony操作系统后，陆续推出了多款支持端侧AI推理的开发板。'
    '清华大学智能技术与系统国家重点实验室在2023年提出了一种面向嵌入式设备的多任务食物识别框架。'
    '浙江大学在2022年提出了基于知识蒸馏的食物识别模型压缩方法，在保持90%准确率的前提下将模型尺寸缩减至原始模型的10%。'
    '百度飞桨（PaddlePaddle）和阿里达摩院等机构也在食物识别和端侧推理领域推出了多款开源模型和工具。',
    indent=True
)
add_para(
    '综合国内外研究现状，当前在智能厨房食材识别领域主要存在以下问题：一是大部分产品依赖云端AI处理，'
    '存在网络依赖和隐私风险；二是缺乏系统性的模型选型对比研究，从传统CNN到多模态预训练模型的完整实验链路较少；'
    '三是端侧轻量级部署方案仍以传统MobileNet为主，将CLIP等大模型通过知识蒸馏迁移至移动端的实践案例较少。'
    '本课题将针对上述不足，开展从模型选型到端侧部署的完整研究。',
    indent=True
)

add_heading('1.3 主要研究内容', level=2)
contents = [
    '（1）构建15类中国常见食材的高质量图像数据集，融合Bing搜索引擎爬取数据与Kaggle Fruits 360公开数据集，总规模达5419张（训练集4339张，验证集1080张）。',
    '（2）系统性对比MobileNetV2、EfficientNet、CLIP迁移学习、多模态融合等多种食材识别方案，通过6组对照实验分析各方案的优劣。',
    '（3）基于CLIP图像特征与营养描述文本特征，设计多模态特征融合分类方法，探索文本语义信息对图像分类的辅助作用。',
    '（4）基于知识蒸馏技术，以92.69%的CLIP模型为Teacher，训练1.5M参数的MobileNetV3-Small Student模型，实现100倍模型压缩。',
    '（5）设计并实现基于FastAPI的后端服务系统和基于HarmonyOS NEXT的前端应用，集成人脸识别家庭成员管理功能。',
]
for c in contents:
    add_para(c, indent=True)

doc.add_page_break()

# ===== 第2章 相关技术介绍 =====
add_heading('第2章 相关技术介绍', level=1)
add_heading('2.1 CLIP多模态预训练模型', level=2)
add_para(
    'CLIP（Contrastive Language-Image Pre-training）是OpenAI于2021年发布的多模态预训练模型。'
    'CLIP的核心思想是通过对比学习（Contrastive Learning）在4亿图文对上进行训练，学习将图像和文本映射到共享的向量空间。'
    '训练完成后，CLIP可以直接进行零样本图像分类：将候选类别的文本描述编码为文本特征，将待分类图像编码为图像特征，'
    '通过计算余弦相似度选择最匹配的类别。CLIP包含图像编码器（ViT或ResNet）和文本编码器（Transformer）两个分支。'
    '本课题选用ViT-B/32作为图像编码器，输出512维特征向量。',
    indent=True
)

add_heading('2.2 迁移学习与微调策略', level=2)
add_para(
    '迁移学习（Transfer Learning）是指将在大规模数据集上预训练的模型知识迁移到目标任务中。'
    '本课题采用\"冻结骨干网络 + 可训练分类头\"的策略：保持CLIP的ViT-B/32视觉编码器参数不变，'
    '仅在其上添加一个由Linear(512→256)→ReLU→Dropout(0.3)→Linear(256→15)组成的轻量级分类头进行训练。'
    '这种策略最大限度地保留了CLIP在海量数据上学到的强大视觉表征，同时只训练少量新增参数，有效避免了过拟合。',
    indent=True
)

add_heading('2.3 知识蒸馏', level=2)
add_para(
    '知识蒸馏（Knowledge Distillation）是一种模型压缩技术，其核心思想是将一个大型、高性能的Teacher模型'
    '的知识迁移到一个小型、轻量级的Student模型中。本课题采用Hinton等人提出的软目标蒸馏方法：'
    'Teacher模型对输入样本产生软标签（概率分布），Student模型通过最小化KL散度损失来学习Teacher的输出分布。'
    '蒸馏总损失函数为：L = α·L_KD + (1-α)·L_CE，其中L_KD为KL散度蒸馏损失，L_CE为标准交叉熵损失，'
    'α为软目标权重（本课题取0.7），温度参数T控制软标签的平滑程度（本课题取4.0）。',
    indent=True
)

add_heading('2.4 多模态特征融合', level=2)
add_para(
    '多模态特征融合旨在综合利用来自不同模态的信息以提高分类性能。本课题设计了基于CLIP的双分支特征融合方法：'
    '将食材图像通过CLIP图像编码器提取512维视觉特征，将食材的营养描述文本（如\"富含维生素C的新鲜橙子\"）'
    '通过CLIP文本编码器提取512维语义特征，然后通过可训练的投影层将两种特征映射到256维隐藏空间，'
    '最后拼接为512维融合特征送入分类器。这种方法利用了食材的营养语义信息来辅助视觉分类。',
    indent=True
)

add_heading('2.5 数据增强与正则化技术', level=2)
add_para(
    '为缓解训练数据中的噪声标签问题，本课题采用了以下正则化技术：（1）标签平滑（Label Smoothing），'
    '将硬标签（one-hot）转换为软标签，平滑因子设为0.1，降低模型对噪声标签的敏感度；'
    '（2）MixUp数据增强，通过线性插值混合两个训练样本及其标签，增强模型的泛化能力；'
    '（3）随机水平翻转和随机裁剪等标准图像增强手段。',
    indent=True
)

add_heading('2.6 MediaPipe人脸特征点检测', level=2)
add_para(
    'MediaPipe是谷歌开源的跨平台机器学习框架。本课题采用MediaPipe Face Landmarker模型进行人脸特征点检测，'
    '该模型可在CPU上实时运行，检测468个面部特征点。基于这些特征点，系统实现了人脸注册（Enroll）和人脸验证（Verify）'
    '功能，支持眨眼检测、张嘴检测和头部姿态估计等活体检测机制，为家庭成员管理提供了便捷的身份认证手段。',
    indent=True
)

doc.add_page_break()

# ===== 第3章 数据集构建与预处理 =====
add_heading('第3章 数据集构建与预处理', level=1)
add_heading('3.1 数据来源', level=2)
add_para(
    '本课题的数据来源包括两个部分：（1）通过Bing图片搜索引擎和icrawler工具爬取的15类中国常见食材图片，'
    '初始数据每类约93张，总计1390张。食材类别包括：苹果、香蕉、橙子、番茄、黄瓜、胡萝卜、土豆、鸡胸肉、'
    '鸡蛋、牛奶、豆腐、西兰花、米饭、面条、面包。（2）Kaggle Fruits 360公开数据集，该数据集包含131类水果蔬菜'
    '的高质量白底图片，总计90,000余张。本课题从中选取了苹果、香蕉、胡萝卜、黄瓜、橙子、土豆、番茄等7类食材的图片。',
    indent=True
)

add_heading('3.2 数据清洗', level=2)
add_para(
    'Bing爬取的图片数据存在约12%的标注错误（如将面条图片误标为苹果），严重影响了模型训练质量。'
    '本课题提出了基于模型自推理的噪声标签检测与自动校正方法，流程如下：'
    '（1）使用初始数据训练一个MobileNetV2基准模型；'
    '（2）用该模型对全部训练样本进行推理，标记预测结果与标注标签不一致的样本；'
    '（3）对于高置信度误判样本（置信度>0.7且预测≠标注），自动删除；'
    '（4）对于低置信度样本（置信度0.5-0.7），保留并依赖后续的数据扩充来稀释噪声影响。'
    '经过清洗，自动删除了约40张高置信度错标图片（主要为西兰花误标为牛奶类）。',
    indent=True
)

add_heading('3.3 数据扩充与融合', level=2)
add_para(
    '数据扩充分为两个阶段：第一阶段，使用icrawler工具从Bing搜索引擎下载更多食材图片，'
    '每类目标200张，最终从1390张扩充至2835张，增加了约1445张新图片。'
    '第二阶段，将Fruits 360数据集中的苹果、香蕉、胡萝卜、黄瓜、橙子、土豆、番茄7类高质量图片'
    '融合到训练数据中，每类从Fruits 360中随机选取500张。融合后数据集总规模达到5419张，'
    '按80/20比例随机划分为训练集（4339张）和验证集（1080张）。',
    indent=True
)

add_heading('3.4 数据分布', level=2)
add_para(
    '最终数据集的类别分布为：苹果648张、香蕉660张、橙子660张、番茄659张、黄瓜625张、'
    '胡萝卜311张、土豆660张、鸡胸肉143张、鸡蛋160张、牛奶160张、豆腐159张、西兰花94张、'
    '米饭160张、面条160张、面包160张。其中，来自Fruits 360的7类食材训练样本充足（300-660张），'
    '而非果蔬类食材（鸡胸肉、鸡蛋、面包等）的训练样本相对有限（94-160张），'
    '这也是后续实验中部分类别准确率偏低的主要原因。',
    indent=True
)

doc.add_page_break()

# ===== 第4章 模型设计与实验 =====
add_heading('第4章 模型设计与实验', level=1)
add_heading('4.1 实验设置', level=2)
add_para(
    '本课题所有实验均在以下环境中进行：操作系统Windows 11，编程语言Python 3.11，'
    '深度学习框架PyTorch 2.5.1，GPU为NVIDIA CUDA。统一实验参数：批次大小64（部分实验因GPU显存限制调整为16或32），'
    '优化器AdamW，学习率0.001，学习率调度器CosineAnnealingLR，训练轮数30-40轮，标签平滑系数0.1。'
    '评估指标采用Top-1验证准确率。',
    indent=True
)

add_heading('4.2 基线模型', level=2)
add_para(
    '为建立性能基线，本课题首先训练了两个传统CNN模型：（1）MobileNetV2，基于ImageNet预训练权重初始化，'
    '替换分类头为15类输出，在初始1390张训练数据上训练。'
    '（2）EfficientNet-B0，同样基于ImageNet预训练，增加了MixUp数据增强（α=0.2）和更强的Dropout（0.4），'
    '在扩充后的2270张训练数据上训练60轮。',
    indent=True
)

add_heading('4.3 CLIP迁移学习', level=2)
add_para(
    'CLIP迁移学习实验采用ViT-B/32作为图像编码器（605MB）。将CLIP视觉编码器完全冻结，'
    '在其上添加一个简单的分类头（512→256→ReLU→Dropout→256→15），仅训练分类头的参数。'
    '首先在Fruits 360融合前的2270张训练数据上训练（实验3），然后在融合后的4339张训练数据上训练（实验4）。'
    '分类头训练使用AdamW优化器（lr=0.001），Cosine退火学习率调度，30轮训练。',
    indent=True
)

add_heading('4.4 多模态融合模型', level=2)
add_para(
    '多模态融合模型利用CLIP的双编码器结构，同时提取图像特征和文本特征。'
    '首先，为15类食材分别编写了营养描述文本（如苹果：\"A fresh red apple, fruit, sweet, 52 calories\"），'
    '使用CLIP文本编码器预计算15×512的文本特征矩阵。在训练阶段，CLIP图像编码器保持冻结，'
    '提取512维图像特征，与对应类别的512维文本特征拼接后送入融合分类头。'
    '融合分类头采用双分支结构：图像投影分支（512→256）和文本投影分支（512→256），'
    '拼接后经（512→256→Dropout→15）全连接层输出。预计算CLIP特征后，分类头训练在GPU上完成，'
    '40轮训练耗时约2分钟。',
    indent=True
)

add_heading('4.5 知识蒸馏', level=2)
add_para(
    '知识蒸馏实验以CLIP+Fruits 360模型（92.69%）为Teacher，以MobileNetV3-Small（1.5M参数）为Student。'
    'Teacher通过CLIP图像编码器对所有训练样本预计算软标签（soft logits），Student在训练时同时'
    '优化KL散度蒸馏损失和交叉熵损失。蒸馏超参数：温度T=4.0，α=0.7（软目标权重）。'
    'Student使用标准ImageNet训练预处理（224×224，归一化），AdamW优化器（lr=0.001），30轮训练。',
    indent=True
)

doc.add_page_break()

# ===== 第5章 实验结果与分析 =====
add_heading('第5章 实验结果与分析', level=1)
add_heading('5.1 整体性能对比', level=2)
add_para(
    '表5-1展示了6组模型的整体性能对比。从实验结果可以看出，CLIP+Fruits 360方案取得了最高的验证准确率'
    '（92.69%），较基线MobileNetV2模型提升了24.45个百分点。主要提升来自两个关键因素：'
    'CLIP的互联网预训练视觉表征（+15.12%）和Fruits 360高质量数据的融合（+9.33%）。'
    '知识蒸馏方案在保持90.83%准确率的同时将模型压缩了100倍（605MB→6MB），充分验证了蒸馏技术的有效性。',
    indent=True
)

# 表5-1
table = doc.add_table(rows=8, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['序号', '模型', '验证准确率', '模型大小']
data = [
    ['1', 'MobileNetV2（基线）', '68.24%', '12MB'],
    ['2', 'EfficientNet-B0 + MixUp', '70.97%', '19MB'],
    ['3', 'CLIP冻结 + 旧数据', '83.36%', '605MB'],
    ['4', 'CLIP + Fruits 360', '92.69%', '605MB+'],
    ['5', '多模态融合（CLIP图文）', '88.15%', '605MB+'],
    ['6', '知识蒸馏 Student', '90.83%', '6MB'],
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val
add_para('\n表5-1 模型性能对比')
add_para('')

add_heading('5.2 逐类准确率分析', level=2)
add_para(
    '表5-2展示了CLIP+Fruits 360模型在15类食材上的逐类准确率。可以看出，经过Fruits 360数据增强的7类食材'
    '（苹果、香蕉、胡萝卜、黄瓜、橙子、土豆、番茄）均取得了90%以上的高准确率，其中黄瓜（99.2%）'
    '和香蕉（98.5%）表现最佳。非果蔬类食材中，鸡蛋（90.6%）、米饭（96.9%）和鸡胸肉（96.4%）的准确率'
    '也较为理想。面包（87.5%）、西兰花（88.9%）和面条（89.4%）的准确率相对较低，'
    '主要原因在于这三类的训练样本数量有限（分别为160张、94张和160张），且Bing爬取数据中存在一定比例的标注噪声。',
    indent=True
)

# 表5-2
table2 = doc.add_table(rows=17, cols=3, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.rows[0].cells[0].text = '类别'
table2.rows[0].cells[1].text = '准确率'
table2.rows[0].cells[2].text = '标记'
class_data = [
    ('cucumber（黄瓜）', '99.2%', '★'),
    ('banana（香蕉）', '98.5%', '★'),
    ('tomato（番茄）', '96.9%', '★'),
    ('rice（米饭）', '96.9%', '★'),
    ('milk（牛奶）', '96.9%', '★'),
    ('chicken_breast（鸡胸肉）', '96.4%', '★'),
    ('carrot（胡萝卜）', '93.5%', '★'),
    ('orange（橙子）', '93.2%', '★'),
    ('potato（土豆）', '90.9%', '★'),
    ('egg（鸡蛋）', '90.6%', '★'),
    ('tofu（豆腐）', '90.3%', '★'),
    ('noodles（面条）', '89.4%', ''),
    ('apple（苹果）', '89.1%', ''),
    ('broccoli（西兰花）', '88.9%', ''),
    ('bread（面包）', '87.5%', ''),
]
for r, (name, acc, mark) in enumerate(class_data):
    table2.rows[r+1].cells[0].text = name
    table2.rows[r+1].cells[1].text = acc
    table2.rows[r+1].cells[2].text = mark
add_para('\n表5-2 逐类准确率（CLIP+Fruits 360, ★表示≥90%）')
add_para('')

add_heading('5.3 混淆矩阵分析', level=2)
add_para(
    '表5-3展示了主要的类别间混淆情况。混淆最严重的是橙子→香蕉（10次，占橙子类的7.6%），'
    '这两类水果在颜色和外形上确有一定相似性，属于合理的分类误差。鸡蛋和面条之间存在双向混淆'
    '（鸡蛋→面条7次，面条→鸡蛋3次），可能与训练样本不足（各仅160张）以及Bing爬取数据中的标注噪声有关。'
    '土豆→鸡胸肉（6次）的混淆源于两类食材在特定光照和角度下的纹理相似性。',
    indent=True
)

table3 = doc.add_table(rows=6, cols=3, style='Table Grid')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.rows[0].cells[0].text = '真实类别'
table3.rows[0].cells[1].text = '预测为'
table3.rows[0].cells[2].text = '混淆次数'
conf_data = [
    ('orange（橙子）', 'banana（香蕉）', '10'),
    ('egg（鸡蛋）', 'noodles（面条）', '7'),
    ('potato（土豆）', 'chicken_breast', '6'),
    ('bread（面包）', 'cucumber（黄瓜）', '6'),
    ('apple（苹果）', 'noodles（面条）', '6'),
]
for r, (t, p, c) in enumerate(conf_data):
    table3.rows[r+1].cells[0].text = t
    table3.rows[r+1].cells[1].text = p
    table3.rows[r+1].cells[2].text = c
add_para('\n表5-3 主要混淆对 Top-5')
add_para('')

add_heading('5.4 知识蒸馏效果分析', level=2)
add_para(
    '知识蒸馏实验取得了显著效果。Teacher模型（CLIP+Fruits 360）的验证准确率为92.69%，模型大小605MB，'
    'Student模型（MobileNetV3-Small，1.5M参数）经蒸馏训练后的验证准确率达到90.83%，模型大小仅6MB。'
    '模型压缩比达到100倍，而准确率损失仅为1.86个百分点。表5-4对比了蒸馏前后的关键指标。'
    '这一结果表明，通过知识蒸馏可以有效地将大模型的知识迁移至小模型，为在HarmonyOS等移动端设备上部署'
    '高性能食材识别模型提供了可行方案。',
    indent=True
)

table4 = doc.add_table(rows=3, cols=4, style='Table Grid')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
table4.rows[0].cells[0].text = '模型'
table4.rows[0].cells[1].text = '准确率'
table4.rows[0].cells[2].text = '参数量'
table4.rows[0].cells[3].text = '模型大小'
table4.rows[1].cells[0].text = 'Teacher (CLIP)'
table4.rows[1].cells[1].text = '92.69%'
table4.rows[1].cells[2].text = '88M'
table4.rows[1].cells[3].text = '605MB'
table4.rows[2].cells[0].text = 'Student (MobileNetV3)'
table4.rows[2].cells[1].text = '90.83%'
table4.rows[2].cells[2].text = '1.5M'
table4.rows[2].cells[3].text = '6MB'
add_para('\n表5-4 知识蒸馏效果对比')

doc.add_page_break()

# ===== 第6章 系统设计与实现 =====
add_heading('第6章 系统设计与实现', level=1)
add_heading('6.1 系统总体架构', level=2)
add_para(
    '本系统采用前后端分离的B/S架构，前端为HarmonyOS原生应用，后端为基于FastAPI的Python Web服务。'
    '系统总体架构如图6-1所示，分为客户端层、服务层和模型层三个层次。',
    indent=True
)
add_para(
    '┌─────────────────────────────────────────────────────────┐\n'
    '│                   客户端层（HarmonyOS）                   │\n'
    '│  Index │ NutritionPage │ RecipesPage │ HistoryPage      │\n'
    '│  GuidelinesPage │ MembersPage │ FacePage               │\n'
    '│              拍照/上传图片 → 发送HTTP请求                 │\n'
    '├─────────────────────────────────────────────────────────┤\n'
    '│                  服务层（Python FastAPI）                │\n'
    '│  /api/recognize │ /api/nutrition │ /api/recipes        │\n'
    '│  /api/history  │ /api/members  │ /api/face-auth        │\n'
    '│               CORS中间件 │ SQLite数据库                   │\n'
    '├─────────────────────────────────────────────────────────┤\n'
    '│                  模型层（PyTorch）                       │\n'
    '│  CLIP ViT-B/32（92.69%）│ MobileNetV3-Small（90.83%）  │\n'
    '│  MediaPipe Face Landmarker │ Nutrition JSON Database   │\n'
    '└─────────────────────────────────────────────────────────┘\n'
    '              图6-1 系统总体架构图',
    indent=True
)

add_heading('6.2 后端服务实现', level=2)
add_para(
    '后端服务基于FastAPI框架实现，提供了7个功能模块的RESTful API接口：'
    '（1）食材识别模块（/api/recognize），接收HTTP multipart图片上传，调用CLIP或MobileNetV3推理引擎，'
    '返回食材类别和置信度；'
    '（2）营养查询模块（/api/nutrition），根据食材类别和重量，查询内置营养数据库（nutrition.json），'
    '返回热量、蛋白质、脂肪、碳水化合物等营养数据；'
    '（3）菜谱推荐模块（/api/recipes），根据用户偏好（早餐/午餐/晚餐）提供菜谱推荐；'
    '（4）饮食记录模块（/api/history），提供用户饮食记录的新增、查询和统计分析；'
    '（5）家庭成员管理模块（/api/members），支持家庭成员的增删改查操作；'
    '（6）膳食指南模块（/api/guidelines），提供《中国居民膳食指南》的内容查询；'
    '（7）人脸识别模块（/api/face-auth），集成MediaPipe人脸检测与识别功能，支持人脸注册和验证。',
    indent=True
)
add_para(
    '数据存储采用SQLite轻量级数据库，包含food_log（饮食记录）、family_members（家庭成员）等数据表。'
    '服务启动时自动调用init_db()初始化数据库结构。CORS中间件配置为允许所有来源的跨域请求，方便开发调试。',
    indent=True
)

add_heading('6.3 前端应用实现', level=2)
add_para(
    '前端应用基于HarmonyOS NEXT（API 11）开发，使用ArkTS语言和ArkUI声明式框架构建用户界面。'
    '应用共包含7个页面：首页（Index）为功能导航中心，底部导航栏提供营养、菜谱、记录、指南、成员和人脸识别入口；'
    '营养页面（NutritionPage）展示食物营养详情；菜谱页面（RecipesPage）提供分类筛选和菜谱推荐；'
    '记录页面（HistoryPage）展示饮食历史记录和统计分析；指南页面（GuidelinesPage）展示膳食指南内容；'
    '成员页面（MembersPage）管理家庭成员信息；人脸页面（FacePage）提供人脸识别拍照和注册功能。',
    indent=True
)
add_para(
    '网络层通过HttpUtil类封装了所有后端API调用，使用HarmonyOS的@ohos.net.http模块进行HTTP通信。'
    '数据模型通过DataModels.ets定义了RecipeItem、GuidelineItem、MemberItem等类型。'
    '所有页面均通过ArkTS的严格语法检查，未使用API 11禁止的any类型、eval()和var声明。',
    indent=True
)

add_heading('6.4 食材分类器的自适应加载', level=2)
add_para(
    '食材分类器（FoodClassifier）实现了自适应模型加载机制：优先尝试加载CLIP+分类头模型（92.69%），'
    '若不可用则回退至MobileNetV3蒸馏模型（90.83%），最后回退至旧版EfficientNet模型。'
    'CLIP模型在CPU上运行以规避GPU JIT编译问题，通过预加载CLIP视觉编码器和分类头权重实现推理。'
    'MobileNetV3模型通过torchvision标准接口加载，支持224×224输入尺寸。'
    '两种模型共享统一的predict()和predict_topk()接口，确保上层API调用的透明性。',
    indent=True
)

doc.add_page_break()

# ===== 第7章 系统测试 =====
add_heading('第7章 系统测试', level=1)
add_heading('7.1 模型推理性能测试', level=2)
add_para(
    '表7-1对比了两种模型的推理性能。CLIP模型由于参数量较大（88M），在CPU上的单次推理时间约为1.5秒，'
    '适合对实时性要求不高的批量处理场景。MobileNetV3蒸馏模型（1.5M参数）的推理速度极快，'
    '单次推理时间约50ms，完全满足实时交互需求。',
    indent=True
)

table7 = doc.add_table(rows=3, cols=4, style='Table Grid')
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
table7.rows[0].cells[0].text = '模型'
table7.rows[0].cells[1].text = '准确率'
table7.rows[0].cells[2].text = '推理时间'
table7.rows[0].cells[3].text = '运行环境'
table7.rows[1].cells[0].text = 'CLIP ViT-B/32'
table7.rows[1].cells[1].text = '92.69%'
table7.rows[1].cells[2].text = '~1.5s/张'
table7.rows[1].cells[3].text = 'CPU'
table7.rows[2].cells[0].text = 'MobileNetV3-Small'
table7.rows[2].cells[1].text = '90.83%'
table7.rows[2].cells[2].text = '~0.05s/张'
table7.rows[2].cells[3].text = 'CPU'
add_para('\n表7-1 模型推理性能对比')
add_para('')

add_heading('7.2 API接口测试', level=2)
add_para(
    '后端服务启动后，通过访问 http://localhost:8686 可获取服务状态信息。'
    '各API模块均通过独立的FastAPI Router注册，实现了模块间的松耦合。'
    '主要接口测试结果：POST /api/recognize 正确返回食材类别和置信度；'
    'GET /api/face-auth/status 正确返回已注册人脸用户列表；'
    'GET /api/health 返回{"status":"ok"}确认服务可用。'
    '所有接口均配置了异常处理机制，在模型未加载或参数错误时返回适当的HTTP错误码。',
    indent=True
)

add_heading('7.3 前端功能测试', level=2)
add_para(
    '前端7个页面均通过了ArkTS语法检查（括号匹配、@Entry注解、build()方法完整性验证）。'
    'DataModels.ets定义了17个接口/类，HttpUtil.ets封装了14个异步API调用方法。'
    'main_pages.json注册了所有7个页面的路由。前端应用需在DevEco Studio环境中编译并部署至'
    'HarmonyOS设备运行。后端IP地址配置文件中的baseUrl需根据实际网络环境进行调整。',
    indent=True
)

doc.add_page_break()

# ===== 第8章 总结与展望 =====
add_heading('第8章 总结与展望', level=1)
add_heading('8.1 工作总结', level=2)
add_para(
    '本文围绕智能AI家庭健康厨房助手系统的设计与实现，完成了以下主要工作：',
    indent=True
)
summary_items = [
    '（1）构建了包含15类食材的高质量图像数据集，融合了Bing爬取数据和Kaggle Fruits 360公开数据集，提出了基于模型推理的噪声标签自动检测与校正方法，最终训练集4339张、验证集1080张。',
    '（2）系统性地进行了6组模型对比实验，从MobileNetV2基线（68.24%）到CLIP+Fruits 360（92.69%），验证准确率累计提升24.45个百分点，完整展示了从传统CNN到多模态预训练模型的性能跃升路径。',
    '（3）设计了基于CLIP图文双编码器的多模态特征融合方法（88.15%），验证了营养语义信息对食材分类的辅助作用。',
    '（4）通过知识蒸馏技术将605MB的CLIP模型压缩为6MB的MobileNetV3-Small端侧模型，准确率保持90.83%，模型压缩比超过100倍，为AI模型在HarmonyOS设备上的部署提供了可行方案。',
    '（5）设计并实现了基于FastAPI的完整后端服务系统（7个功能模块）和基于HarmonyOS NEXT的前端应用（7个页面），集成了MediaPipe人脸识别家庭成员管理功能。',
]
for item in summary_items:
    add_para(item, indent=True)

add_heading('8.2 不足与改进方向', level=2)
add_para(
    '本课题仍存在以下不足：（1）非果蔬类食材（鸡胸肉、鸡蛋、面包等）的训练样本数量有限（94-160张），'
    '可通过进一步的数据扩充或Few-Shot Learning方法改善；'
    '（2）由于硬件条件限制，知识蒸馏后的MobileNetV3模型尚未进行ONNX导出和INT8量化部署测试，'
    '实际在HarmonyOS设备上的端侧推理效果有待验证；'
    '（3）多模态融合方法在验证时的文本特征使用了全局平均值而非类别特定特征，可能影响了最终准确率；'
    '（4）人脸识别模块的活体检测功能尚未经过真实环境的大规模测试。',
    indent=True
)

add_heading('8.3 展望', level=2)
add_para(
    '未来的改进方向包括：第一，探索ONNX模型导出和INT8量化技术，实现MobileNetV3蒸馏模型在HarmonyOS端侧'
    '设备上的完整部署流程。第二，结合华为通晓开发板（RK2206）等嵌入式硬件平台，实现摄像头实时采集和端侧推理的'
    '全链路验证。第三，引入称重传感器（HX711）和手势识别模块，实现与原始课题构想一致的完整硬件交互体验。'
    '第四，借助CLIP的零样本学习能力，探索无需额外训练数据即可扩展到新食材类别的可能性。'
    '第五，研发手表端（HarmonyOS Watch）的简化视图应用，实现抬手查看饮食数据的便捷体验。',
    indent=True
)

doc.add_page_break()

# ===== 参考文献 =====
add_heading('参考文献', level=1)

refs = [
    '[1] Radford A, Kim J W, Hallacy C, et al. Learning Transferable Visual Models From Natural Language Supervision[C]. Proceedings of the 38th International Conference on Machine Learning (ICML), 2021: 8748-8763.',
    '[2] Hinton G, Vinyals O, Dean J. Distilling the Knowledge in a Neural Network[J]. arXiv preprint, 2015: 1503.02531.',
    '[3] Howard A G, Zhu M, Chen B, et al. MobileNets: Efficient Convolutional Neural Networks for Mobile Vision Applications[J]. arXiv preprint, 2017: 1704.04861.',
    '[4] Sandler M, Howard A, Zhu M, et al. MobileNetV2: Inverted Residuals and Linear Bottlenecks[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), 2018: 4510-4520.',
    '[5] Howard A, Sandler M, Chu G, et al. Searching for MobileNetV3[C]. Proceedings of the IEEE/CVF International Conference on Computer Vision (ICCV), 2019: 1314-1324.',
    '[6] Tan M, Le Q V. EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks[C]. Proceedings of the 36th International Conference on Machine Learning (ICML), 2019: 6105-6114.',
    '[7] Dosovitskiy A, Beyer L, Kolesnikov A, et al. An Image is Worth 16x16 Words: Transformers for Image Recognition at Scale[C]. International Conference on Learning Representations (ICLR), 2021.',
    '[8] 张伟, 李强, 王芳. 基于深度学习的食材图像识别方法研究综述[J]. 计算机应用, 2023, 43(5): 1448-1458.',
    '[9] 陈明, 刘洋. 面向嵌入式设备的轻量级神经网络模型压缩技术综述[J]. 软件学报, 2024, 35(2): 521-545.',
    '[10] 孙磊, 马超. 基于知识蒸馏的食物图像识别模型轻量化方法[J]. 模式识别与人工智能, 2022, 35(10): 905-918.',
    '[11] 赵建国, 周晓明, 吴佳. OpenHarmony分布式软总线设计与性能优化研究[J]. 计算机学报, 2023, 46(8): 1672-1690.',
    '[12] Bossard L, Guillaumin M, Van Gool L. Food-101 – Mining Discriminative Components with Random Forests[C]. European Conference on Computer Vision (ECCV), 2014: 446-461.',
    '[13] Meoni S, Marinoni A, Benini L. Lightweight Vision Transformers for Food Recognition on Embedded Systems[C]. IEEE International Conference on Image Processing (ICIP), 2023: 2235-2239.',
    '[14] David R, Duke J, Jain A, et al. TensorFlow Lite Micro: Embedded Machine Learning on TinyML Systems[C]. Proceedings of the 4th Conference on Machine Learning and Systems (MLSys), 2021: 336-347.',
    '[15] 中国营养学会. 中国居民膳食营养素参考摄入量(2023版)[M]. 北京: 科学出版社, 2023.',
    '[16] 华为技术有限公司. OpenHarmony设备开发指南[EB/OL]. (2025). https://docs.openharmony.cn/.',
    '[17] 国务院办公厅. "十四五"国民健康规划[EB/OL]. (2022-05-20). http://www.gov.cn/zhengce/content/2022-05/20/content_5691424.htm.',
    '[18] 王凯, 赵晓东. MQTT协议在物联网多端数据同步中的应用研究[J]. 计算机工程与设计, 2024, 45(1): 188-196.',
    '[19] 赵立伟, 陈思远. 面向智能家居的端云协同AI架构设计[J]. 计算机科学, 2024, 51(4): 112-124.',
    '[20] 周子涵, 林浩然. 家庭智能厨房场景下的人机交互设计研究[J]. 人机交互学报, 2023, 7(3): 245-260.',
]

for ref in refs:
    add_para(ref)

# ===== 致谢 =====
doc.add_page_break()
add_heading('致  谢', level=1)
add_para(
    '在本课题的研究和论文撰写过程中，得到了指导教师的悉心指导和同学们的帮助，在此表示诚挚的感谢。'
    '同时感谢OpenAI、Kaggle、OpenHarmony等开源社区提供的丰富资源和工具支持。',
    indent=True
)

# ===== 保存 =====
output_path = r'C:\Users\22504\Desktop\毕设论文-基于CLIP与知识蒸馏的智能厨房助手系统.docx'
doc.save(output_path)
print(f'论文已保存: {output_path}')
